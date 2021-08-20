#GUIWiki,.py
import wikipedia


#===เปลี่ยนเป็นภาษาไทย===
wikipedia.set_lang('th')

#python to docx
from docx import Document
def Wiki(keyword,lang='th',):
    
    wikipedia.set_lang(lang)

    data = wikipedia.summary(keyword)

    # page + content บทความทั้งหน้า

    data2 = wikipedia.page(keyword)
    data2 = data2.content

    doc = Document()
    doc.add_heading(keyword,0)

    doc.add_paragraph(data2)
    doc.save(keyword + '.docx')
    print('สร้างไฟล์สำเร็จ')




#===== GUI =====
from tkinter import *
from tkinter import ttk
from tkinter import messagebox #เพื่อสั่งโชว์ error ขึ้นมาถ้ามีคำสั่งผิด

GUI = Tk() #สร้าง pop-up
GUI.title('Wikipedia Program by PAWIT')
GUI.geometry('400x300')

#===Config===
FONT1 = ('Angsana New',15)

#===คำอธิบาย===
L= ttk.Label(GUI, text='ค้นหาบทความ',font = FONT1)
L.pack()

#===Search Entry===
v_search = StringVar() #กล่องสำหรับเก็บ keyword
E1 = ttk.Entry(GUI,textvariable = v_search,font = FONT1,width = 40)
E1.pack(pady = 10)

#===ปุ่มค้นหา===
def Search():
    keyword = v_search.get() # .get() คืิอดึงข้อมูลเข้ามา ซึ่งจะใช้ได้ StringVar() เท่านั้น
    try:
        #ลองค้นหา keyword ถ้าหากได้ ให้ผ่านไป
        language = v_radio.get() # th / en / zh
        Wiki(keyword,language)
        messagebox.showinfo('บันทึกสำเร็จ','ค้นหาข้อความสำเร็จ บันทึกเรียบร้อยแล้ว')
    except:
        #หากรันคำสั่งแล้วมีปัญหา แสดงข้อความแจ้งเตือน
        messagebox.showwarning('Keyword Error','กรุณากรอกคำค้นหาใหม่')
            
    # print(wikipedia.search(keyword)) # keyword อื่นๆที่ใกล้เคียงกับ keyword ทีเราหา
    # result = wikipedia.summary(keyword)
    # print(result)
    # print('\n--------------------------------------')

B1 = ttk.Button(GUI,text='Search',command = Search)
B1.pack(ipadx = 20 , ipady = 10)

# เลือกภาษา
F1 = Frame(GUI) #Frame เปรียบเสมือนกระดาน white board
F1.pack(pady = 10)

v_radio = StringVar() # ช่องเก็บข้อมูลภาษา

RB1 = ttk.Radiobutton(F1,text='ภาษาไทย',variable = v_radio,value='th') #ให้ v_radio = 'th'
RB2 = ttk.Radiobutton(F1,text='อังกฤษ',variable = v_radio,value='en') #ให้ v_radio = 'en'
RB3 = ttk.Radiobutton(F1,text='จีน',variable = v_radio,value='zh') #ให้ v_radio = 'zh'
RB1.invoke() #ให้ Default เป็นภาษาไทย

RB1.grid(row=1,column=0) #วางตำแหน่งแบบ grid ใส่เป็นจำแหน่ง 0,0
RB2.grid(row=1,column=1)
RB3.grid(row=1,column=2)


GUI.mainloop()
