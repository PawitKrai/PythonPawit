#test-docx.py
from docx import Document
import wikipedia



def Wiki(keyword,lang='th',):
    
    wikipedia.set_lang(lang)

    data = wikipedia.summary(keyword)

    # page + content บทความทั้งหน้า

    data2 = wikipedia.page(keyword)
    data2 = data2.content

    doc = Document()
    doc.add_heading(keyword,0)

    doc.add_paragraph(data2)
    doc.save(keyword + '.docx')
    print('สร้างไฟล์สำเร็จ')

try:
    Wiki('','en')
except:
    print('ERROR')


#Wiki('สหรัฐอเมริกา','en')
#Wiki('ประเทศญี่ปุ่น','jp')
#Wiki('ประเทศจีน','zh')
